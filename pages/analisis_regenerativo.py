import streamlit as st
import geopandas as gpd
import pandas as pd
import numpy as np
import tempfile
import zipfile
import os
import folium
from streamlit_folium import st_folium
from sentinelhub import SHConfig, DataCollection, MimeType, CRS, BBox, SentinelHubRequest, bbox_to_dimensions
from docx import Document
import io
import base64
from datetime import datetime, timedelta

st.set_page_config(page_title="Análisis Regenerativo", layout="wide")
st.title("Análisis Forrajero Regenerativo")

# --- VERIFICACIÓN RIGUROSA DE CREDENCIALES ---
st.info("Verificando credenciales de Sentinel Hub...")

try:
    # LEER SECRETS
    client_id = st.secrets.get("SENTINEL_HUB_CLIENT_ID", "").strip()
    client_secret = st.secrets.get("SENTINEL_HUB_CLIENT_SECRET", "").strip()
    instance_id = st.secrets.get("SENTINEL_HUB_INSTANCE_ID", "").strip()

    if not client_id or not client_secret:
        raise ValueError("Faltan CLIENT_ID o CLIENT_SECRET")

    # CONFIGURAR
    config = SHConfig()
    config.sh_client_id = client_id
    config.sh_client_secret = client_secret
    if instance_id:
        config.instance_id = instance_id

    # PRUEBA DE CONEXIÓN REAL
    test_request = SentinelHubRequest(
        evalscript="return [1];",
        input_data=[SentinelHubRequest.input_data(data_collection=DataCollection.SENTINEL2_L2A)],
        responses=[SentinelHubRequest.output_response('default', MimeType.JSON)],
        bbox=BBox((0, 0, 1, 1), crs=CRS.WGS84),
        size=(1, 1),
        config=config
    )
    test_request.get_data()  # Si falla, lanza error
    st.success("Credenciales válidas")
    SH_OK = True
except Exception as e:
    SH_OK = False
    st.error(f"Credenciales inválidas: {e}")
    st.code("""
SENTINEL_HUB_CLIENT_ID = "358474d6-2326-4637-bf8e-30a709b2d6a6"
SENTINEL_HUB_CLIENT_SECRET = "b296cf70-c9d2-4e69-91f4-f7be80b99ed1"
SENTINEL_HUB_INSTANCE_ID = "PLAK81593ed161694ad48faa8065411d2539"
    """)
    st.stop()

# --- SIDEBAR ---
with st.sidebar:
    st.header("Parámetros")
    tipo_pastura = st.selectbox("Pastura", ["Alfalfa", "Raygrass", "Festuca", "Natural"])
    default_date = datetime.now().date()
    fecha = st.date_input("Fecha imagen", value=default_date, max_value=default_date)
    nubes = st.slider("Nubes máx (%)", 0, 100, 20)
    peso_vaca = st.slider("Peso promedio (kg)", 400, 600, 500)
    eficiencia = st.slider("Eficiencia pastoreo (%)", 40, 80, 55) / 100

# --- CARGA SHAPEFILE ---
uploaded = st.file_uploader("Subir ZIP con shapefile", type="zip")

if uploaded:
    with tempfile.TemporaryDirectory() as tmp:
        with zipfile.ZipFile(uploaded) as z:
            z.extractall(tmp)
        shp = [f for f in os.listdir(tmp) if f.endswith('.shp')][0]
        gdf = gpd.read_file(f"{tmp}/{shp}")
        if gdf.crs is None:
            gdf = gdf.set_crs('EPSG:4326')
        gdf = gdf.to_crs('EPSG:4326')
        gdf = gdf[gdf.geometry.notna() & gdf.geometry.is_valid]
        st.success(f"{len(gdf)} lotes válidos")
else:
    gdf = None

# --- ANÁLISIS ---
if gdf is not None and st.button("EJECUTAR ANÁLISIS", type="primary"):
    progress = st.progress(0)
    status = st.empty()
    results = []

    for idx, (i, row) in enumerate(gdf.iterrows()):
        status.text(f"Procesando lote {idx+1}/{len(gdf)}...")
        progress.progress((idx + 1) / len(gdf))

        geom = row.geometry
        if geom is None or geom.is_empty:
            ndvi_mean = 0.1
        else:
            bbox = BBox(geom.bounds, crs=CRS.WGS84)
            size = bbox_to_dimensions(bbox, resolution=10)

            evalscript = """
            //VERSION=3
            function setup() { return { input: ["B04", "B08", "dataMask"], output: { bands: 4 } }; }
            function evaluatePixel(s) { let ndvi = (s.B08 - s.B04)/(s.B08 + s.B04 + 0.0001); return [ndvi, ndvi, ndvi, s.dataMask]; }
            """

            request = SentinelHubRequest(
                evalscript=evalscript,
                input_data=[SentinelHubRequest.input_data(
                    data_collection=DataCollection.SENTINEL2_L2A,
                    time_interval=(str(fecha), str(fecha + timedelta(days=1))),
                    other_args={"maxcc": nubes / 100}
                )],
                responses=[SentinelHubRequest.output_response('default', MimeType.TIFF)],
                bbox=bbox,
                size=size,
                config=config
            )

            try:
                response = request.get_data()[0]
                mask = response[..., 3] == 1
                ndvi_values = response[..., 0][mask]
                ndvi_mean = np.mean(ndvi_values) if len(ndvi_values) > 0 else 0.1
            except Exception as e:
                st.warning(f"Lote {idx+1}: {e}")
                ndvi_mean = 0.1

        area_ha = geom.area / 10000 if geom else 0
        biomasa = max(ndvi_mean * 4000, 1000) if ndvi_mean > 0.3 else 1000
        consumo_dia = peso_vaca * 0.025 * eficiencia
        ev_ha = biomasa / (consumo_dia * 30) if consumo_dia > 0 else 0
        dias = biomasa / (consumo_dia * ev_ha) if ev_ha > 0 else 0

        results.append({
            'id': i,
            'area_ha': round(area_ha, 2),
            'ndvi': round(ndvi_mean, 3),
            'biomasa_kg_ha': int(biomasa),
            'ev_ha': round(ev_ha, 2),
            'dias': round(dias, 1)
        })

    progress.empty()
    status.empty()

    df = pd.DataFrame(results)
    gdf_result = gdf.merge(df, left_index=True, right_on='id').set_index('id')

    st.success("¡Análisis completado!")

    # --- MAPA ---
    if not gdf_result.empty:
        m = folium.Map(location=[gdf_result.geometry.centroid.y.mean(), gdf_result.geometry.centroid.x.mean()], zoom_start=14)
        for _, r in gdf_result.iterrows():
            if r.geometry is None or r.geometry.is_empty:
                continue
            coords = []
            if r.geometry.geom_type == 'Polygon':
                coords = list(r.geometry.exterior.coords)
            elif r.geometry.geom_type == 'MultiPolygon':
                for poly in r.geometry.geoms:
                    coords.extend(list(poly.exterior.coords))
            if coords:
                color = "green" if r.ev_ha > 1.5 else "orange" if r.ev_ha > 0.8 else "red"
                folium.Polygon(
                    locations=[(y, x) for x, y in coords],
                    popup=f"EV/ha: {r.ev_ha} | Días: {r.dias} | Biomasa: {r.biomasa_kg_ha} kg/ha",
                    color=color, fill=True, weight=2
                ).add_to(m)
        st_folium(m, width=700, height=500)

    # --- RESULTADOS ---
    st.subheader("Resultados")
    st.dataframe(df.style.format({"area_ha": "{:.2f}", "ndvi": "{:.3f}", "ev_ha": "{:.2f}"}))

    col1, col2 = st.columns(2)
    with col1:
        st.download_button("CSV", df.to_csv(index=False), "resultados.csv", "text/csv")
    with col2:
        st.download_button("GeoJSON", gdf_result.to_json(), "lotes.geojson", "application/json")

    # --- RECOMENDACIONES ---
    st.subheader("Recomendaciones")
    prom_ev = df['ev_ha'].mean()
    prom_dias = df['dias'].mean()
    rec = [
        f"**EV/ha promedio**: {prom_ev:.2f}",
        f"**Días promedio**: {prom_dias:.1f}",
        "**Rotación intensiva**: 1-3 días + 30-60 descanso",
        "**Monitoreo mensual**"
    ]
    for r in rec:
        st.markdown(f"- {r}")

    if st.button("Generar Informe DOCX"):
        doc = Document()
        doc.add_heading('Informe Forrajero', 0)
        doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
        doc.add_paragraph(f"Lotes: {len(gdf_result)} | Área: {df.area_ha.sum():.1f} ha")
        for r in rec:
            doc.add_paragraph(r, style='List Bullet')
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        b64 = base64.b64encode(buf.read()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="informe.docx">Descargar</a>'
        st.markdown(href, unsafe_allow_html=True)
        st.success("Informe listo.")
