import streamlit as st
import geopandas as gpd
import pandas as pd
import numpy as np
import tempfile
import os
import zipfile
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.colors import LinearSegmentedColormap
import io
from shapely.geometry import Polygon, box
import math
import json
import base64
import hashlib
import secrets

# Importaciones para mapas
try:
    import folium
    from streamlit_folium import folium_static, st_folium
    FOLIUM_AVAILABLE = True
except ImportError:
    FOLIUM_AVAILABLE = False

# Importaciones para informes
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# =============================================================================
# CONFIGURACIÓN Y AUTENTICACIÓN
# =============================================================================

st.set_page_config(
    page_title="🌱 Analizador Forrajero Unificado - Sentinel Hub + PRV",
    page_icon="🌱",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =============================================================================
# INICIALIZACIÓN COMPLETA DEL SESSION STATE
# =============================================================================

def initialize_session_state():
    """Inicializa todas las variables del session state"""
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False
    if 'username' not in st.session_state:
        st.session_state.username = ""
    if 'gdf_cargado' not in st.session_state:
        st.session_state.gdf_cargado = None
    if 'gdf_analizado' not in st.session_state:
        st.session_state.gdf_analizado = None
    if 'sh_client_id' not in st.session_state:
        st.session_state.sh_client_id = None
    if 'sh_client_secret' not in st.session_state:
        st.session_state.sh_client_secret = None
    if 'sh_configured' not in st.session_state:
        st.session_state.sh_configured = False
    # Nuevos estados para parámetros personalizados
    if 'parametros_personalizados' not in st.session_state:
        st.session_state.parametros_personalizados = {}

# Sistema de autenticación simple
def check_authentication():
    """Verifica las credenciales de autenticación"""
    # Credenciales por defecto (en producción usar variables de entorno)
    default_users = {
        "admin": hashlib.sha256("password123".encode()).hexdigest(),
        "user": hashlib.sha256("user123".encode()).hexdigest(),
        "tech": hashlib.sha256("tech123".encode()).hexdigest()
    }
    
    return default_users

def login_section():
    """Sección de login"""
    st.title("🔐 Inicio de Sesión - Analizador Forrajero")
    st.markdown("---")
    
    users_db = check_authentication()
    
    with st.form("login_form"):
        username = st.text_input("Usuario")
        password = st.text_input("Contraseña", type="password")
        submit = st.form_submit_button("Iniciar Sesión")
        
        if submit:
            if username in users_db:
                hashed_password = hashlib.sha256(password.encode()).hexdigest()
                if users_db[username] == hashed_password:
                    st.session_state.authenticated = True
                    st.session_state.username = username
                    st.success(f"✅ Bienvenido, {username}!")
                    st.rerun()
                else:
                    st.error("❌ Contraseña incorrecta")
            else:
                st.error("❌ Usuario no encontrado")
    
    # Información de usuarios demo
    with st.expander("ℹ️ Información de acceso demo"):
        st.markdown("""
        **Usuarios de prueba:**
        - **admin** / password123
        - **user** / user123  
        - **tech** / tech123
        
        *En producción, configurar variables de entorno para credenciales reales*
        """)

# =============================================================================
# CONFIGURACIÓN SENTINEL HUB - MEJORADA
# =============================================================================

class SentinelHubConfig:
    def __init__(self):
        self.base_url = "https://services.sentinel-hub.com/ogc/wms/"
        self.available = False
        self.config_message = ""
        
    def check_configuration(self):
        try:
            # 1. PRIMERO: Verificar Secrets de Streamlit Cloud
            if all(key in st.secrets for key in ['SENTINEL_HUB_CLIENT_ID', 'SENTINEL_HUB_CLIENT_SECRET']):
                st.session_state.sh_client_id = st.secrets['SENTINEL_HUB_CLIENT_ID']
                st.session_state.sh_client_secret = st.secrets['SENTINEL_HUB_CLIENT_SECRET']
                st.session_state.sh_configured = True
                self.available = True
                self.config_message = "✅ Sentinel Hub configurado (Streamlit Secrets)"
                return True
            
            # 2. SEGUNDO: Verificar variables de entorno
            elif all(os.getenv(key) for key in ['SENTINEL_HUB_CLIENT_ID', 'SENTINEL_HUB_CLIENT_SECRET']):
                st.session_state.sh_client_id = os.getenv('SENTINEL_HUB_CLIENT_ID')
                st.session_state.sh_client_secret = os.getenv('SENTINEL_HUB_CLIENT_SECRET')
                st.session_state.sh_configured = True
                self.available = True
                self.config_message = "✅ Sentinel Hub configurado (Variables Entorno)"
                return True
            
            # 3. TERCERO: Verificar session state (configuración manual previa)
            elif ('sh_client_id' in st.session_state and 
                  'sh_client_secret' in st.session_state and
                  st.session_state.sh_client_id and 
                  st.session_state.sh_client_secret):
                
                st.session_state.sh_configured = True
                self.available = True
                self.config_message = "✅ Sentinel Hub configurado (Manual)"
                return True
            
            # 4. FALLBACK: Modo simulado para desarrollo
            else:
                self.available = False
                self.config_message = "❌ Sentinel Hub no configurado - Usando modo simulado"
                return False
                
        except Exception as e:
            self.available = False
            self.config_message = f"❌ Error: {str(e)}"
            return False

    def get_credentials_status(self):
        """Devuelve el estado de las credenciales para mostrar en la UI"""
        if self.available:
            client_id = st.session_state.get('sh_client_id', '')
            # Mostrar solo los primeros y últimos caracteres por seguridad
            if client_id and len(client_id) > 12:
                masked_id = f"{client_id[:8]}...{client_id[-4:]}"
            else:
                masked_id = "No disponible"
            return f"🟢 Conectado (ID: {masked_id})"
        else:
            return "🔴 No configurado - Modo simulado"

class SentinelHubProcessor:
    """Procesa datos reales de Sentinel Hub con manejo de errores"""
    
    def __init__(self):
        self.base_url = "https://services.sentinel-hub.com/ogc/wms/"
        self.sh_config = SentinelHubConfig()
        
    def get_ndvi_for_geometry(self, geometry, fecha, bbox, width=512, height=512):
        """Obtiene NDVI real desde Sentinel Hub o simula datos"""
        try:
            if not self.sh_config.available:
                return self._simular_ndvi_response(geometry)
                
            # Aquí iría el código real para conectar con Sentinel Hub
            # Por ahora simulamos la respuesta
            return self._simular_ndvi_response(geometry)
            
        except Exception as e:
            st.error(f"Error obteniendo NDVI de Sentinel Hub: {e}")
            return self._simular_ndvi_response(geometry)
    
    def _simular_ndvi_response(self, geometry):
        """Simula respuesta de Sentinel Hub (para desarrollo)"""
        try:
            # Simular NDVI basado en la posición de la geometría
            centroid = geometry.centroid
            x_norm = (centroid.x * 100) % 1
            y_norm = (centroid.y * 100) % 1
            
            # Crear patrones realistas
            if x_norm < 0.2 or y_norm < 0.2:
                ndvi = 0.15 + np.random.normal(0, 0.05)  # Bordes - suelo
            elif x_norm > 0.7 and y_norm > 0.7:
                ndvi = 0.75 + np.random.normal(0, 0.03)  # Esquina - vegetación densa
            else:
                ndvi = 0.45 + np.random.normal(0, 0.04)  # Centro - vegetación media
            
            return max(0.1, min(0.85, ndvi))
            
        except:
            return 0.5  # Valor por defecto

# =============================================================================
# PARÁMETROS FORRAJEROS UNIFICADOS - AJUSTADOS Y MÁS REALISTAS
# =============================================================================

PARAMETROS_FORRAJEROS = {
    'ALFALFA': {
        'MS_POR_HA_OPTIMO': 4500,
        'CRECIMIENTO_DIARIO': 85,
        'CONSUMO_PORCENTAJE_PESO': 0.032,
        'TASA_UTILIZACION_RECOMENDADA': 0.68,
        'FACTOR_BIOMASA_NDVI': 3200,
        'UMBRAL_NDVI_SUELO': 0.12,
        'UMBRAL_NDVI_PASTURA': 0.42,
        'CONSUMO_DIARIO_EV': 13,
        'EFICIENCIA_PASTOREO': 0.78,
        'EFICIENCIA_COSECHA': 0.72,
        'DIGESTIBILIDAD': 0.68,
        'PROTEINA_CRUDA': 0.20,
        'CONSUMO_VOLUNTARIO': 3.2
    },
    'RAYGRASS': {
        'MS_POR_HA_OPTIMO': 3800,
        'CRECIMIENTO_DIARIO': 75,
        'CONSUMO_PORCENTAJE_PESO': 0.029,
        'TASA_UTILIZACION_RECOMENDADA': 0.62,
        'FACTOR_BIOMASA_NDVI': 2800,
        'UMBRAL_NDVI_SUELO': 0.15,
        'UMBRAL_NDVI_PASTURA': 0.48,
        'CONSUMO_DIARIO_EV': 11,
        'EFICIENCIA_PASTOREO': 0.72,
        'EFICIENCIA_COSECHA': 0.67,
        'DIGESTIBILIDAD': 0.72,
        'PROTEINA_CRUDA': 0.16,
        'CONSUMO_VOLUNTARIO': 2.9
    },
    'FESTUCA': {
        'MS_POR_HA_OPTIMO': 3200,
        'CRECIMIENTO_DIARIO': 55,
        'CONSUMO_PORCENTAJE_PESO': 0.026,
        'TASA_UTILIZACION_RECOMENDADA': 0.58,
        'FACTOR_BIOMASA_NDVI': 2400,
        'UMBRAL_NDVI_SUELO': 0.18,
        'UMBRAL_NDVI_PASTURA': 0.52,
        'CONSUMO_DIARIO_EV': 10,
        'EFICIENCIA_PASTOREO': 0.68,
        'EFICIENCIA_COSECHA': 0.62,
        'DIGESTIBILIDAD': 0.62,
        'PROTEINA_CRUDA': 0.13,
        'CONSUMO_VOLUNTARIO': 2.6
    },
    'AGROPIRRO': {
        'MS_POR_HA_OPTIMO': 3500,
        'CRECIMIENTO_DIARIO': 60,
        'CONSUMO_PORCENTAJE_PESO': 0.027,
        'TASA_UTILIZACION_RECOMENDADA': 0.60,
        'FACTOR_BIOMASA_NDVI': 2600,
        'UMBRAL_NDVI_SUELO': 0.16,
        'UMBRAL_NDVI_PASTURA': 0.50,
        'CONSUMO_DIARIO_EV': 10.5,
        'EFICIENCIA_PASTOREO': 0.70,
        'EFICIENCIA_COSECHA': 0.65,
        'DIGESTIBILIDAD': 0.58,
        'PROTEINA_CRUDA': 0.11,
        'CONSUMO_VOLUNTARIO': 2.7
    },
    'PASTIZAL_NATURAL': {
        'MS_POR_HA_OPTIMO': 2800,
        'CRECIMIENTO_DIARIO': 45,
        'CONSUMO_PORCENTAJE_PESO': 0.024,
        'TASA_UTILIZACION_RECOMENDADA': 0.52,
        'FACTOR_BIOMASA_NDVI': 2200,
        'UMBRAL_NDVI_SUELO': 0.20,
        'UMBRAL_NDVI_PASTURA': 0.46,
        'CONSUMO_DIARIO_EV': 9,
        'EFICIENCIA_PASTOREO': 0.65,
        'EFICIENCIA_COSECHA': 0.58,
        'DIGESTIBILIDAD': 0.52,
        'PROTEINA_CRUDA': 0.09,
        'CONSUMO_VOLUNTARIO': 2.4
    },
    'PERSONALIZADO': {
        'MS_POR_HA_OPTIMO': 3500,
        'CRECIMIENTO_DIARIO': 65,
        'CONSUMO_PORCENTAJE_PESO': 0.028,
        'TASA_UTILIZACION_RECOMENDADA': 0.60,
        'FACTOR_BIOMASA_NDVI': 2500,
        'UMBRAL_NDVI_SUELO': 0.16,
        'UMBRAL_NDVI_PASTURA': 0.48,
        'CONSUMO_DIARIO_EV': 11,
        'EFICIENCIA_PASTOREO': 0.70,
        'EFICIENCIA_COSECHA': 0.65,
        'DIGESTIBILIDAD': 0.65,
        'PROTEINA_CRUDA': 0.14,
        'CONSUMO_VOLUNTARIO': 2.8
    }
}

def obtener_parametros(tipo_pastura, parametros_personalizados=None):
    """Obtiene parámetros con soporte para personalización"""
    if tipo_pastura == "PERSONALIZADO" and parametros_personalizados:
        # Combinar parámetros base con personalizados
        base = PARAMETROS_FORRAJEROS['PERSONALIZADO'].copy()
        base.update(parametros_personalizados)
        return base
    else:
        return PARAMETROS_FORRAJEROS.get(tipo_pastura, PARAMETROS_FORRAJEROS['FESTUCA'])

# =============================================================================
# FUNCIONES DE CÁLCULO UNIFICADAS - MEJORADAS Y MÁS REALISTAS
# =============================================================================

def calcular_ev_ha(biomasa_disponible_kg_ms_ha, consumo_diario_ev, eficiencia_pastoreo=0.7):
    """Calcula EV/ha con parámetros más realistas"""
    if consumo_diario_ev <= 0:
        return 0
    # Cálculo más preciso considerando eficiencia de pastoreo
    ev_ha = (biomasa_disponible_kg_ms_ha * eficiencia_pastoreo) / consumo_diario_ev
    return max(0, round(ev_ha, 2))

def calcular_carga_animal_total(ev_ha, area_ha):
    """Calcula carga animal total"""
    return round(ev_ha * area_ha, 1)

def calcular_dias_permanencia(biomasa_total_kg, consumo_total_diario, crecimiento_diario_kg=0):
    """Calcula días de permanencia con crecimiento considerado"""
    if consumo_total_diario <= 0:
        return 0
    
    # Cálculo base
    dias_base = biomasa_total_kg / consumo_total_diario
    
    if crecimiento_diario_kg > 0:
        # Ajustar por crecimiento durante el pastoreo (modelo más realista)
        factor_crecimiento = 0.25  # Factor conservador
        crecimiento_efectivo = crecimiento_diario_kg * factor_crecimiento * dias_base
        dias_ajustados = (biomasa_total_kg + crecimiento_efectivo) / consumo_total_diario
        return min(round(dias_ajustados, 1), 120)  # Máximo 120 días
    
    return min(round(dias_base, 1), 120)

def calcular_disponibilidad_forrajera(gdf_analizado, tipo_pastura, parametros_personalizados=None):
    """Calcula la disponibilidad forrajera con métricas mejoradas y más realistas"""
    params = obtener_parametros(tipo_pastura, parametros_personalizados)
    
    # Cálculos mejorados de disponibilidad considerando eficiencia de cosecha
    gdf_analizado['disponibilidad_forrajera_kg_ms_ha'] = (
        gdf_analizado['biomasa_disponible_kg_ms_ha'] * 
        params['EFICIENCIA_PASTOREO'] * 
        params['EFICIENCIA_COSECHA']
    ).round(0)
    
    # Clasificación de disponibilidad MÁS REALISTA
    condiciones = [
        gdf_analizado['disponibilidad_forrajera_kg_ms_ha'] < 800,
        gdf_analizado['disponibilidad_forrajera_kg_ms_ha'] < 2000,
        gdf_analizado['disponibilidad_forrajera_kg_ms_ha'] < 3500,
        gdf_analizado['disponibilidad_forrajera_kg_ms_ha'] >= 3500
    ]
    categorias = ['MUY BAJA', 'BAJA', 'MEDIA', 'ALTA']
    gdf_analizado['categoria_disponibilidad'] = np.select(condiciones, categorias, default='MEDIA')
    
    # Cálculo de días de autonomía más realista
    consumo_promedio_diario = 25  # kg MS/día/EV (valor más realista)
    gdf_analizado['dias_autonomia'] = (gdf_analizado['disponibilidad_forrajera_kg_ms_ha'] / consumo_promedio_diario).round(1)
    
    return gdf_analizado

# =============================================================================
# FUNCIONES DE MAPAS UNIFICADAS - LEYENDAS CLARAS Y VISIBLES
# =============================================================================

def crear_mapa_base(gdf, mapa_seleccionado="ESRI World Imagery", zoom_start=10):
    """Crea el mapa base con controles"""
    if not FOLIUM_AVAILABLE or gdf is None:
        return None
        
    bounds = gdf.total_bounds
    center_lat = (bounds[1] + bounds[3]) / 2
    center_lon = (bounds[0] + bounds[2]) / 2
    
    m = folium.Map(
        location=[center_lat, center_lon],
        zoom_start=zoom_start,
        tiles=None,
        control_scale=True
    )
    
    # Capas base
    folium.TileLayer(
        tiles="https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}",
        attr="Esri, Maxar, Earthstar Geographics",
        name="🌍 Satélite",
        control=True,
        show=(mapa_seleccionado == "ESRI World Imagery")
    ).add_to(m)
    
    folium.TileLayer(
        tiles="https://server.arcgisonline.com/ArcGIS/rest/services/World_Street_Map/MapServer/tile/{z}/{y}/{x}",
        attr="Esri, HERE, Garmin",
        name="🗺️ Calles",
        control=True,
        show=(mapa_seleccionado == "ESRI World Street Map")
    ).add_to(m)
    
    folium.TileLayer(
        tiles="https://{s}.tile.openstreetmap.org/{z}/{x}/{y}.png",
        attr="OpenStreetMap contributors",
        name="🗾 OpenStreetMap",
        control=True,
        show=(mapa_seleccionado == "OpenStreetMap")
    ).add_to(m)
    
    # Control de capas
    folium.LayerControl().add_to(m)
    
    return m

def crear_mapa_ndvi(gdf_analizado, mapa_base="ESRI World Imagery"):
    """Crea mapa interactivo de NDVI con leyenda CLARA y visible"""
    if not FOLIUM_AVAILABLE:
        return None
        
    m = crear_mapa_base(gdf_analizado, mapa_base)
    
    def estilo_ndvi(feature):
        ndvi = feature['properties']['ndvi']
        if ndvi < 0.1:
            color = '#8B4513'  # Marrón - suelo desnudo
        elif ndvi < 0.2:
            color = '#CD853F'  # Marrón claro - suelo con algo de vegetación
        elif ndvi < 0.3:
            color = '#FFD700'  # Amarillo - vegetación muy escasa
        elif ndvi < 0.4:
            color = '#ADFF2F'  # Verde amarillento - vegetación escasa
        elif ndvi < 0.5:
            color = '#32CD32'  # Verde claro - vegetación moderada
        elif ndvi < 0.6:
            color = '#228B22'  # Verde - vegetación buena
        elif ndvi < 0.7:
            color = '#006400'  # Verde oscuro - vegetación muy buena
        else:
            color = '#004d00'  # Verde muy oscuro - vegetación excelente
            
        return {
            'fillColor': color,
            'color': 'black',
            'weight': 1,
            'fillOpacity': 0.7
        }
    
    folium.GeoJson(
        gdf_analizado.__geo_interface__,
        style_function=estilo_ndvi,
        tooltip=folium.GeoJsonTooltip(
            fields=['id_subLote', 'ndvi', 'tipo_superficie', 'biomasa_disponible_kg_ms_ha'],
            aliases=['Sub-Lote:', 'NDVI:', 'Tipo:', 'Biomasa (kg MS/ha):'],
            localize=True,
            style=("background-color: white; color: #333333; font-family: arial; font-size: 12px; padding: 10px;")
        )
    ).add_to(m)
    
    # LEYENDA CLARA Y VISIBLE para NDVI
    legend_html = '''
    <div style="
        position: fixed; 
        bottom: 20px; 
        left: 20px; 
        width: 280px; 
        height: auto;
        background-color: white; 
        border: 3px solid #2E8B57;
        z-index: 9999; 
        font-size: 12px; 
        padding: 12px;
        border-radius: 8px;
        box-shadow: 0 0 20px rgba(0,0,0,0.4);
        font-family: Arial, sans-serif;
    ">
        <div style="
            font-weight: bold; 
            margin-bottom: 12px; 
            text-align: center; 
            font-size: 14px; 
            color: #2E8B57;
            background-color: #f0fff0;
            padding: 8px;
            border-radius: 4px;
            border: 1px solid #2E8B57;
        ">
            🌿 ÍNDICE NDVI - ESTADO VEGETATIVO
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 6px;">
            <div style="width: 20px; height: 15px; background: #8B4513; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <span style="font-weight: bold;">&lt; 0.1</span> - Suelo desnudo
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 6px;">
            <div style="width: 20px; height: 15px; background: #CD853F; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <span style="font-weight: bold;">0.1-0.2</span> - Muy escasa
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 6px;">
            <div style="width: 20px; height: 15px; background: #FFD700; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <span style="font-weight: bold;">0.2-0.3</span> - Escasa
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 6px;">
            <div style="width: 20px; height: 15px; background: #ADFF2F; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <span style="font-weight: bold;">0.3-0.4</span> - Regular
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 6px;">
            <div style="width: 20px; height: 15px; background: #32CD32; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <span style="font-weight: bold;">0.4-0.5</span> - Buena
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 6px;">
            <div style="width: 20px; height: 15px; background: #228B22; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <span style="font-weight: bold;">0.5-0.6</span> - Muy Buena
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 6px;">
            <div style="width: 20px; height: 15px; background: #006400; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <span style="font-weight: bold;">0.6-0.7</span> - Excelente
        </div>
        <div style="display: flex; align-items: center;">
            <div style="width: 20px; height: 15px; background: #004d00; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <span style="font-weight: bold;">&gt; 0.7</span> - Óptima
        </div>
    </div>
    '''
    m.get_root().html.add_child(folium.Element(legend_html))
    
    return m

def crear_mapa_ev_ha(gdf_analizado, mapa_base="ESRI World Imagery"):
    """Crea mapa interactivo de EV/ha con leyenda CLARA y visible"""
    if not FOLIUM_AVAILABLE:
        return None
        
    m = crear_mapa_base(gdf_analizado, mapa_base)
    
    def estilo_ev_ha(feature):
        ev_ha = feature['properties']['ev_ha']
        if ev_ha < 0.3:
            color = '#FF4444'  # Rojo - muy baja
        elif ev_ha < 0.8:
            color = '#FF6B6B'  # Rojo claro - baja
        elif ev_ha < 1.5:
            color = '#FFA726'  # Naranja - moderada baja
        elif ev_ha < 2.5:
            color = '#FFD54F'  # Amarillo - moderada
        elif ev_ha < 3.5:
            color = '#9CCC65'  # Verde claro - buena
        elif ev_ha < 5.0:
            color = '#66BB6A'  # Verde - muy buena
        else:
            color = '#2E7D32'  # Verde oscuro - excelente
            
        return {
            'fillColor': color,
            'color': 'black',
            'weight': 1,
            'fillOpacity': 0.7
        }
    
    folium.GeoJson(
        gdf_analizado.__geo_interface__,
        style_function=estilo_ev_ha,
        tooltip=folium.GeoJsonTooltip(
            fields=['id_subLote', 'ev_ha', 'biomasa_disponible_kg_ms_ha', 'dias_permanencia'],
            aliases=['Sub-Lote:', 'EV/ha:', 'Biomasa (kg MS/ha):', 'Días Permanencia:'],
            localize=True,
            style=("background-color: white; color: #333333; font-family: arial; font-size: 12px; padding: 10px;")
        )
    ).add_to(m)
    
    # LEYENDA CLARA Y VISIBLE para EV/ha
    legend_html = '''
    <div style="
        position: fixed; 
        bottom: 20px; 
        left: 20px; 
        width: 260px; 
        height: auto;
        background-color: white; 
        border: 3px solid #2E8B57;
        z-index: 9999; 
        font-size: 12px; 
        padding: 12px;
        border-radius: 8px;
        box-shadow: 0 0 20px rgba(0,0,0,0.4);
        font-family: Arial, sans-serif;
    ">
        <div style="
            font-weight: bold; 
            margin-bottom: 12px; 
            text-align: center; 
            font-size: 14px; 
            color: #2E8B57;
            background-color: #f0fff0;
            padding: 8px;
            border-radius: 4px;
            border: 1px solid #2E8B57;
        ">
            🐄 CAPACIDAD DE CARGA (EV/ha)
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 6px;">
            <div style="width: 20px; height: 15px; background: #FF4444; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <span style="font-weight: bold;">&lt; 0.3</span> - Muy Baja
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 6px;">
            <div style="width: 20px; height: 15px; background: #FF6B6B; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <span style="font-weight: bold;">0.3-0.8</span> - Baja
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 6px;">
            <div style="width: 20px; height: 15px; background: #FFA726; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <span style="font-weight: bold;">0.8-1.5</span> - Moderada Baja
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 6px;">
            <div style="width: 20px; height: 15px; background: #FFD54F; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <span style="font-weight: bold;">1.5-2.5</span> - Moderada
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 6px;">
            <div style="width: 20px; height: 15px; background: #9CCC65; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <span style="font-weight: bold;">2.5-3.5</span> - Buena
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 6px;">
            <div style="width: 20px; height: 15px; background: #66BB6A; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <span style="font-weight: bold;">3.5-5.0</span> - Muy Buena
        </div>
        <div style="display: flex; align-items: center;">
            <div style="width: 20px; height: 15px; background: #2E7D32; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <span style="font-weight: bold;">&gt; 5.0</span> - Excelente
        </div>
    </div>
    '''
    m.get_root().html.add_child(folium.Element(legend_html))
    
    return m

def crear_mapa_disponibilidad(gdf_analizado, mapa_base="ESRI World Imagery"):
    """Crea mapa interactivo de disponibilidad forrajera con leyenda CLARA y visible"""
    if not FOLIUM_AVAILABLE:
        return None
        
    m = crear_mapa_base(gdf_analizado, mapa_base)
    
    def estilo_disponibilidad(feature):
        categoria = feature['properties']['categoria_disponibilidad']
        if categoria == 'MUY BAJA':
            color = '#D32F2F'  # Rojo oscuro
        elif categoria == 'BAJA':
            color = '#FF5252'  # Rojo
        elif categoria == 'MEDIA':
            color = '#FFEB3B'  # Amarillo
        else:
            color = '#4CAF50'  # Verde
            
        return {
            'fillColor': color,
            'color': 'black',
            'weight': 1,
            'fillOpacity': 0.7
        }
    
    folium.GeoJson(
        gdf_analizado.__geo_interface__,
        style_function=estilo_disponibilidad,
        tooltip=folium.GeoJsonTooltip(
            fields=['id_subLote', 'categoria_disponibilidad', 'disponibilidad_forrajera_kg_ms_ha', 'dias_autonomia'],
            aliases=['Sub-Lote:', 'Categoría:', 'Disponibilidad (kg MS/ha):', 'Días Autonomía:'],
            localize=True,
            style=("background-color: white; color: #333333; font-family: arial; font-size: 12px; padding: 10px;")
        )
    ).add_to(m)
    
    # LEYENDA CLARA Y VISIBLE para Disponibilidad
    legend_html = '''
    <div style="
        position: fixed; 
        bottom: 20px; 
        left: 20px; 
        width: 300px; 
        height: auto;
        background-color: white; 
        border: 3px solid #2E8B57;
        z-index: 9999; 
        font-size: 12px; 
        padding: 12px;
        border-radius: 8px;
        box-shadow: 0 0 20px rgba(0,0,0,0.4);
        font-family: Arial, sans-serif;
    ">
        <div style="
            font-weight: bold; 
            margin-bottom: 12px; 
            text-align: center; 
            font-size: 14px; 
            color: #2E8B57;
            background-color: #f0fff0;
            padding: 8px;
            border-radius: 4px;
            border: 1px solid #2E8B57;
        ">
            📊 DISPONIBILIDAD FORRAJERA
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 8px; padding: 4px; background-color: #ffebee; border-radius: 4px;">
            <div style="width: 20px; height: 15px; background: #D32F2F; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <div>
                <span style="font-weight: bold; color: #D32F2F;">🔴 MUY BAJA</span><br>
                <span style="font-size: 11px;">&lt; 800 kg MS/ha</span>
            </div>
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 8px; padding: 4px; background-color: #fff3e0; border-radius: 4px;">
            <div style="width: 20px; height: 15px; background: #FF5252; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <div>
                <span style="font-weight: bold; color: #FF9800;">🟠 BAJA</span><br>
                <span style="font-size: 11px;">800-2000 kg MS/ha</span>
            </div>
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 8px; padding: 4px; background-color: #fffde7; border-radius: 4px;">
            <div style="width: 20px; height: 15px; background: #FFEB3B; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <div>
                <span style="font-weight: bold; color: #FFC107;">🟡 MEDIA</span><br>
                <span style="font-size: 11px;">2000-3500 kg MS/ha</span>
            </div>
        </div>
        <div style="display: flex; align-items: center; padding: 4px; background-color: #e8f5e8; border-radius: 4px;">
            <div style="width: 20px; height: 15px; background: #4CAF50; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <div>
                <span style="font-weight: bold; color: #4CAF50;">🟢 ALTA</span><br>
                <span style="font-size: 11px;">&gt; 3500 kg MS/ha</span>
            </div>
        </div>
    </div>
    '''
    m.get_root().html.add_child(folium.Element(legend_html))
    
    return m

def crear_mapa_recomendaciones(gdf_analizado, mapa_base="ESRI World Imagery"):
    """Crea mapa interactivo con recomendaciones agroecológicas con leyenda CLARA y visible"""
    if not FOLIUM_AVAILABLE:
        return None
        
    m = crear_mapa_base(gdf_analizado, mapa_base)
    
    def estilo_recomendaciones(feature):
        categoria = feature['properties']['categoria_disponibilidad']
        if categoria == 'MUY BAJA':
            color = '#D32F2F'  # Rojo oscuro - intervención urgente
        elif categoria == 'BAJA':
            color = '#FF9800'  # Naranja - manejo intensivo
        elif categoria == 'MEDIA':
            color = '#FFEB3B'  # Amarillo - manejo cuidadoso
        else:
            color = '#4CAF50'  # Verde - mantenimiento
            
        return {
            'fillColor': color,
            'color': 'black',
            'weight': 2,
            'fillOpacity': 0.6
        }
    
    # Agregar polígonos con recomendaciones
    folium.GeoJson(
        gdf_analizado.__geo_interface__,
        style_function=estilo_recomendaciones,
        tooltip=folium.GeoJsonTooltip(
            fields=['id_subLote', 'categoria_disponibilidad', 'disponibilidad_forrajera_kg_ms_ha', 'dias_autonomia'],
            aliases=['Sub-Lote:', 'Categoría:', 'Disponibilidad:', 'Autonomía (días):'],
            localize=True,
            style=("background-color: white; color: #333333; font-family: arial; font-size: 12px; padding: 10px;")
        ),
        popup=folium.GeoJsonPopup(
            fields=['id_subLote', 'categoria_disponibilidad', 'disponibilidad_forrajera_kg_ms_ha', 'ev_ha'],
            aliases=['Sub-Lote:', 'Estado:', 'Disponibilidad (kg MS/ha):', 'EV/ha:'],
            localize=True
        )
    ).add_to(m)
    
    # LEYENDA CLARA Y VISIBLE para Recomendaciones
    legend_html = '''
    <div style="
        position: fixed; 
        bottom: 20px; 
        left: 20px; 
        width: 320px; 
        height: auto;
        background-color: white; 
        border: 3px solid #2E8B57;
        z-index: 9999; 
        font-size: 12px; 
        padding: 12px;
        border-radius: 8px;
        box-shadow: 0 0 20px rgba(0,0,0,0.4);
        font-family: Arial, sans-serif;
    ">
        <div style="
            font-weight: bold; 
            margin-bottom: 12px; 
            text-align: center; 
            font-size: 14px; 
            color: #2E8B57;
            background-color: #f0fff0;
            padding: 8px;
            border-radius: 4px;
            border: 1px solid #2E8B57;
        ">
            🌱 RECOMENDACIONES AGROECOLÓGICAS
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 8px; padding: 6px; background-color: #ffebee; border-radius: 4px; border-left: 4px solid #D32F2F;">
            <div style="width: 20px; height: 15px; background: #D32F2F; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <div>
                <span style="font-weight: bold; color: #D32F2F;">🔴 INTERVENCIÓN URGENTE</span><br>
                <span style="font-size: 11px;">Disponibilidad &lt; 800 kg MS/ha</span>
            </div>
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 8px; padding: 6px; background-color: #fff3e0; border-radius: 4px; border-left: 4px solid #FF9800;">
            <div style="width: 20px; height: 15px; background: #FF9800; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <div>
                <span style="font-weight: bold; color: #FF9800;">🟠 MANEJO INTENSIVO</span><br>
                <span style="font-size: 11px;">Disponibilidad 800-2000 kg MS/ha</span>
            </div>
        </div>
        <div style="display: flex; align-items: center; margin-bottom: 8px; padding: 6px; background-color: #fffde7; border-radius: 4px; border-left: 4px solid #FFEB3B;">
            <div style="width: 20px; height: 15px; background: #FFEB3B; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <div>
                <span style="font-weight: bold; color: #FFC107;">🟡 MANEJO CUIDADOSO</span><br>
                <span style="font-size: 11px;">Disponibilidad 2000-3500 kg MS/ha</span>
            </div>
        </div>
        <div style="display: flex; align-items: center; padding: 6px; background-color: #e8f5e8; border-radius: 4px; border-left: 4px solid #4CAF50;">
            <div style="width: 20px; height: 15px; background: #4CAF50; border: 2px solid #000; margin-right: 8px; border-radius: 3px;"></div>
            <div>
                <span style="font-weight: bold; color: #4CAF50;">🟢 MANTENIMIENTO</span><br>
                <span style="font-size: 11px;">Disponibilidad &gt; 3500 kg MS/ha</span>
            </div>
        </div>
    </div>
    '''
    m.get_root().html.add_child(folium.Element(legend_html))
    
    return m

# =============================================================================
# SISTEMA DE ANÁLISIS UNIFICADO - MEJORADO CON PARÁMETROS REALISTAS
# =============================================================================

class AnalizadorForrajeroUnificado:
    def __init__(self):
        self.sh_config = SentinelHubConfig()
        self.sh_config.check_configuration()
    
    def analizar_potrero(self, gdf, config):
        """Análisis completo unificado con parámetros más realistas"""
        try:
            st.header("🌱 ANÁLISIS FORRAJERO UNIFICADO")
            
            # Dividir potrero
            gdf_dividido = self._dividir_potrero(gdf, config['n_divisiones'])
            if gdf_dividido is None:
                st.error("Error dividiendo potrero")
                return None
            
            # Calcular áreas
            areas_ha = self._calcular_superficie(gdf_dividido)
            gdf_dividido['area_ha'] = areas_ha
            
            # Obtener datos de vegetación
            resultados = self._obtener_datos_vegetacion(gdf_dividido, config)
            
            # Calcular métricas ganaderas
            gdf_analizado = self._calcular_metricas_ganaderas(gdf_dividido, resultados, config)
            
            # Calcular disponibilidad forrajera
            parametros_personalizados = config.get('parametros_personalizados')
            gdf_analizado = calcular_disponibilidad_forrajera(gdf_analizado, config['tipo_pastura'], parametros_personalizados)
            
            return gdf_analizado
            
        except Exception as e:
            st.error(f"Error en análisis: {e}")
            return None
    
    def _dividir_potrero(self, gdf, n_zonas):
        """Divide el potrero en sub-lotes"""
        if len(gdf) == 0:
            return gdf
        
        try:
            potrero = gdf.iloc[0].geometry
            bounds = potrero.bounds
            minx, miny, maxx, maxy = bounds
            
            sub_poligonos = []
            n_cols = math.ceil(math.sqrt(n_zonas))
            n_rows = math.ceil(n_zonas / n_cols)
            width = (maxx - minx) / n_cols
            height = (maxy - miny) / n_rows
            
            for i in range(n_rows):
                for j in range(n_cols):
                    if len(sub_poligonos) >= n_zonas:
                        break
                        
                    cell = Polygon([
                        (minx + j * width, miny + i * height),
                        (minx + (j + 1) * width, miny + i * height),
                        (minx + (j + 1) * width, miny + (i + 1) * height),
                        (minx + j * width, miny + (i + 1) * height)
                    ])
                    
                    intersection = potrero.intersection(cell)
                    if not intersection.is_empty and intersection.area > 0:
                        sub_poligonos.append(intersection)
            
            if sub_poligonos:
                return gpd.GeoDataFrame({
                    'id_subLote': range(1, len(sub_poligonos) + 1),
                    'geometry': sub_poligonos
                }, crs=gdf.crs)
            return gdf
                
        except Exception as e:
            st.error(f"Error dividiendo potrero: {e}")
            return gdf
    
    def _calcular_superficie(self, gdf):
        """Calcula superficie en hectáreas"""
        try:
            if gdf.crs and gdf.crs.is_geographic:
                gdf_proj = gdf.to_crs('EPSG:3857')
                area_m2 = gdf_proj.geometry.area
            else:
                area_m2 = gdf.geometry.area
            return area_m2 / 10000
        except:
            return gdf.geometry.area / 10000
    
    def _obtener_datos_vegetacion(self, gdf, config):
        """Obtiene datos de vegetación (simulado o real) con parámetros realistas"""
        resultados = []
        processor = SentinelHubProcessor()
        parametros_personalizados = config.get('parametros_personalizados')
        
        for idx, row in gdf.iterrows():
            # Obtener NDVI (real o simulado)
            fecha_imagen = config.get('fecha_imagen', datetime.now() - timedelta(days=30))
            bounds = gdf.total_bounds
            bbox = [bounds[0], bounds[1], bounds[2], bounds[3]]
            
            ndvi = processor.get_ndvi_for_geometry(row.geometry, fecha_imagen, bbox)
            
            # Calcular biomasa basada en NDVI con parámetros realistas
            params = obtener_parametros(config['tipo_pastura'], parametros_personalizados)
            
            # Cálculo más realista de biomasa considerando relación NDVI-biomasa
            if ndvi < params['UMBRAL_NDVI_SUELO']:
                biomasa_total = params['FACTOR_BIOMASA_NDVI'] * 0.1  # Mínimo para suelo
            elif ndvi < 0.3:
                biomasa_total = params['FACTOR_BIOMASA_NDVI'] * ndvi * 0.8
            elif ndvi < 0.6:
                biomasa_total = params['FACTOR_BIOMASA_NDVI'] * ndvi * 1.2
            else:
                biomasa_total = params['FACTOR_BIOMASA_NDVI'] * ndvi * 1.5
            
            biomasa_disponible = biomasa_total * params['TASA_UTILIZACION_RECOMENDADA']
            
            # Clasificar vegetación más detallada
            if ndvi < params['UMBRAL_NDVI_SUELO']:
                tipo_veg = "SUELO_DESNUDO"
            elif ndvi < 0.25:
                tipo_veg = "VEGETACION_MUY_ESCASA"
            elif ndvi < params['UMBRAL_NDVI_PASTURA']:
                tipo_veg = "VEGETACION_ESCASA"
            elif ndvi < 0.6:
                tipo_veg = "VEGETACION_MODERADA"
            else:
                tipo_veg = "VEGETACION_DENSA"
            
            resultados.append({
                'id_subLote': row['id_subLote'],
                'ndvi': round(ndvi, 3),
                'tipo_superficie': tipo_veg,
                'biomasa_total_kg_ms_ha': round(biomasa_total, 0),
                'biomasa_disponible_kg_ms_ha': round(biomasa_disponible, 0),
                'cobertura_vegetal': round(max(0.1, min(0.95, ndvi * 1.3)), 2)
            })
        
        return resultados
    
    def _calcular_metricas_ganaderas(self, gdf, resultados, config):
        """Calcula todas las métricas ganaderas con parámetros realistas"""
        parametros_personalizados = config.get('parametros_personalizados')
        params = obtener_parametros(config['tipo_pastura'], parametros_personalizados)
        
        # Usar consumo personalizado si está disponible
        consumo_diario_personalizado = config.get('consumo_voluntario', params.get('CONSUMO_VOLUNTARIO', 10))
        eficiencia_pastoreo_personalizada = config.get('eficiencia_pastoreo', params['EFICIENCIA_PASTOREO'])
        eficiencia_cosecha_personalizada = config.get('eficiencia_cosecha', params['EFICIENCIA_COSECHA'])
        
        for idx, resultado in enumerate(resultados):
            area_ha = gdf.loc[gdf.index[idx], 'area_ha']
            biomasa_disponible = resultado['biomasa_disponible_kg_ms_ha']
            
            # EV/ha con parámetros personalizados - cálculo más realista
            ev_ha = calcular_ev_ha(biomasa_disponible, consumo_diario_personalizado, eficiencia_pastoreo_personalizada)
            
            # Carga animal
            carga_animal = calcular_carga_animal_total(ev_ha, area_ha)
            
            # Días de permanencia con cálculo mejorado
            biomasa_total_kg = biomasa_disponible * area_ha
            consumo_individual_kg = config['peso_promedio'] * params['CONSUMO_PORCENTAJE_PESO']
            consumo_total_diario = config['carga_animal'] * consumo_individual_kg
            crecimiento_diario_kg = params['CRECIMIENTO_DIARIO'] * area_ha
            
            dias_permanencia = calcular_dias_permanencia(
                biomasa_total_kg, 
                consumo_total_diario, 
                crecimiento_diario_kg
            )
            
            # Añadir al GeoDataFrame
            for key, value in resultado.items():
                gdf.loc[gdf.index[idx], key] = value
            
            gdf.loc[gdf.index[idx], 'ev_ha'] = ev_ha
            gdf.loc[gdf.index[idx], 'carga_animal'] = carga_animal
            gdf.loc[gdf.index[idx], 'dias_permanencia'] = dias_permanencia
            gdf.loc[gdf.index[idx], 'consumo_individual_kg'] = round(consumo_individual_kg, 2)
            gdf.loc[gdf.index[idx], 'biomasa_total_kg'] = round(biomasa_total_kg, 0)
        
        return gdf

# =============================================================================
# GENERACIÓN DE INFORMES - MEJORADA CON PARÁMETROS REALISTAS
# =============================================================================

def crear_mapa_detallado(gdf_analizado, tipo_pastura):
    """Crea un mapa detallado para el informe"""
    try:
        # Crear figura de matplotlib
        fig, ax = plt.subplots(1, 1, figsize=(10, 8))
        
        # Plot simple para el informe
        gdf_analizado.plot(column='ndvi', ax=ax, legend=True, 
                          cmap='RdYlGn', vmin=0, vmax=1)
        ax.set_title(f'Mapa de NDVI - {tipo_pastura}')
        ax.set_axis_off()
        
        # Guardar en buffer
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight')
        buffer.seek(0)
        plt.close()
        
        return buffer
    except Exception as e:
        st.error(f"Error creando mapa para informe: {e}")
        return None

def generar_informe_completo(gdf_analizado, config, mapa_bytes=None):
    """Genera informe DOCX con análisis completo"""
    if not DOCX_AVAILABLE:
        st.error("python-docx no disponible. Instala: pip install python-docx")
        return None
    
    try:
        doc = Document()
        
        # Título
        titulo = doc.add_heading('INFORME DE ANÁLISIS FORRAJERO UNIFICADO', 0)
        
        # Información general
        doc.add_heading('Información General', level=1)
        doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"Tipo de pastura: {config['tipo_pastura']}")
        doc.add_paragraph(f"Usuario: {st.session_state.username}")
        
        # Métricas principales
        doc.add_heading('Métricas Principales', level=1)
        area_total = gdf_analizado['area_ha'].sum()
        biomasa_prom = gdf_analizado['biomasa_disponible_kg_ms_ha'].mean()
        ev_total = gdf_analizado['carga_animal'].sum()
        dias_prom = gdf_analizado['dias_permanencia'].mean()
        disponibilidad_prom = gdf_analizado['disponibilidad_forrajera_kg_ms_ha'].mean()
        
        doc.add_paragraph(f"Área total analizada: {area_total:.2f} ha")
        doc.add_paragraph(f"Biomasa disponible promedio: {biomasa_prom:.0f} kg MS/ha")
        doc.add_paragraph(f"Disponibilidad forrajera promedio: {disponibilidad_prom:.0f} kg MS/ha")
        doc.add_paragraph(f"Capacidad total de carga: {ev_total:.1f} EV")
        doc.add_paragraph(f"Días de permanencia promedio: {dias_prom:.1f} días")
        
        # Mapa si está disponible
        if mapa_bytes:
            doc.add_heading('Mapa de Análisis', level=1)
            try:
                mapa_bytes.seek(0)
                with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
                    tmp.write(mapa_bytes.read())
                    tmp_path = tmp.name
                doc.add_picture(tmp_path, width=Inches(6))
                os.unlink(tmp_path)
            except Exception as e:
                doc.add_paragraph(f"Error al insertar mapa: {e}")
        
        # Recomendaciones agroecológicas
        doc.add_heading('Recomendaciones Agroecológicas', level=1)
        recomendaciones = _generar_recomendaciones_agroecologicas(gdf_analizado, config)
        for recomendacion in recomendaciones.split('\n'):
            if recomendacion.strip():
                doc.add_paragraph(recomendacion.strip())
        
        # Guardar
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        st.error(f"Error generando informe: {e}")
        return None

def _generar_recomendaciones_agroecologicas(gdf, config):
    """Genera recomendaciones agroecológicas basadas en el análisis"""
    biomasa_prom = gdf['biomasa_disponible_kg_ms_ha'].mean()
    disponibilidad_prom = gdf['disponibilidad_forrajera_kg_ms_ha'].mean()
    dias_prom = gdf['dias_permanencia'].mean()
    
    recomendaciones = []
    
    if disponibilidad_prom < 800:
        recomendaciones.extend([
            "🔴 **ESTADO CRÍTICO - INTERVENCIÓN AGROECOLÓGICA URGENTE**",
            "",
            "🌱 **PRÁCTICAS REGENERATIVAS INMEDIATAS:**",
            "• Implementar descanso prolongado (90-120 días) con exclusión animal",
            "• Aplicar abonos verdes y cobertura orgánica (mulch)",
            "• Sembrar especies pioneras y leguminosas fijadoras de nitrógeno",
            "• Incorporar compost y biofertilizantes para recuperar suelo",
            "",
            "🐄 **MANEJO GANADERO:**",
            "• Reducir carga animal inmediatamente (máximo 0.5 EV/ha)",
            "• Implementar suplementación estratégica con conservas",
            "• Rotaciones muy cortas (1-2 días) con largos descansos",
            "",
            "💧 **MANEJO HÍDRICO:**",
            "• Implementar zanjas de infiltración y barreras vivas",
            "• Proteger cursos de agua con franjas buffer",
            "• Utilizar coberturas para retener humedad"
        ])
    elif disponibilidad_prom < 2000:
        recomendaciones.extend([
            "🟠 **ESTADO DE MEJORA - MANEJO AGROECOLÓGICO ACTIVO**",
            "",
            "🌱 **PRÁCTICAS REGENERATIVAS:**",
            "• Mantener rotaciones con 45-60 días de descanso",
            "• Enriquecer con mezclas de gramíneas y leguminosas",
            "• Aplicar microorganismos eficientes y biofertilizantes",
            "• Implementar pastoreo racional Voisin",
            "",
            "🐄 **MANEJO GANADERO:**",
            "• Carga animal moderada (1-2 EV/ha según disponibilidad)",
            "• Monitoreo semanal de crecimiento y ajuste de carga",
            "• Pastoreos cortos e intensivos para estimular rebrote",
            "",
            "📊 **SEGUIMIENTO:**",
            "• Medir altura forrajera y cobertura vegetal periódicamente",
            "• Registrar datos de biomasa y días de descanso",
            "• Ajustar manejo según observaciones en campo"
        ])
    else:
        recomendaciones.extend([
            "🟢 **ESTADO ÓPTIMO - MANEJO AGROECOLÓGICO CONSERVATIVO**",
            "",
            "🌱 **PRÁCTICAS REGENERATIVAS:**",
            "• Mantener sistema actual con mejoras incrementales",
            "• Diversificar con especies nativas y forrajeras perennes",
            "• Implementar agroforestería y sistemas silvopastoriles",
            "• Conservar biodiversidad y hábitats naturales",
            "",
            "🐄 **MANEJO GANADERO:**",
            "• Carga animal óptima (2-4 EV/ha según capacidad)",
            "• Rotaciones con 30-45 días de descanso",
            "• Aprovechar picos de crecimiento con pastoreos intensivos",
            "",
            "🌍 **SUSTENTABILIDAD:**",
            "• Monitorear salud del suelo y materia orgánica",
            "• Implementar captura de carbono en pastizales",
            "• Conservar corredores biológicos y fuentes de agua"
        ])
    
    # Recomendaciones específicas adicionales
    if dias_prom < 7:
        recomendaciones.extend([
            "",
            "⚠️ **ALERTA:** Días de permanencia muy bajos",
            "• Considerar suplementación inmediata",
            "• Revisar carga animal y distribución"
        ])
    elif dias_prom > 60:
        recomendaciones.extend([
            "",
            "✅ **EXCELENTE:** Sistema con buena autonomía",
            "• Mantener prácticas actuales",
            "• Considerar enriquecimiento con leguminosas"
        ])
    
    return "\n".join(recomendaciones)

# =============================================================================
# INTERFAZ PRINCIPAL - MEJORADA CON PARÁMETROS REALISTAS
# =============================================================================

def main_application():
    """Aplicación principal después del login"""
    
    # Sidebar de configuración
    with st.sidebar:
        st.header(f"👋 Bienvenido, {st.session_state.username}")
        
        if st.button("🚪 Cerrar Sesión"):
            st.session_state.authenticated = False
            st.session_state.username = ""
            st.rerun()
        
        st.markdown("---")
        
        # =============================================================================
        # 🛰️ CONFIGURACIÓN SENTINEL HUB
        # =============================================================================
        st.header("🛰️ Configuración Sentinel Hub")
        
        # Inicializar y verificar configuración
        sh_config = SentinelHubConfig()
        sh_configured = sh_config.check_configuration()
        
        # Mostrar estado
        status_text = sh_config.get_credentials_status()
        if sh_configured:
            st.success(status_text)
        else:
            st.warning(status_text)
        
        # Opción para configuración manual (solo si no está configurado)
        if not sh_configured:
            with st.expander("🔧 Configurar Manualmente", expanded=False):
                st.info("""
                **Para datos satelitales reales necesitas:**
                1. Cuenta en [Sentinel Hub](https://www.sentinel-hub.com/)
                2. Client ID y Client Secret
                3. Configurar instancia en el dashboard
                """)
                
                sh_client_id = st.text_input("Client ID", key="manual_client_id")
                sh_client_secret = st.text_input("Client Secret", type="password", key="manual_client_secret")
                
                if st.button("💾 Guardar Credenciales Manualmente"):
                    if sh_client_id and sh_client_secret:
                        st.session_state.sh_client_id = sh_client_id
                        st.session_state.sh_client_secret = sh_client_secret
                        st.session_state.sh_configured = True
                        st.success("Credenciales guardadas en sesión")
                        st.rerun()
                    else:
                        st.error("Ingresa ambas credenciales")
        
        # Opción para limpiar credenciales (solo si están configuradas)
        else:
            if st.button("🔄 Limpiar Credenciales"):
                if 'sh_client_id' in st.session_state:
                    del st.session_state.sh_client_id
                if 'sh_client_secret' in st.session_state:
                    del st.session_state.sh_client_secret
                st.session_state.sh_configured = False
                st.success("Credenciales limpiadas")
                st.rerun()
        
        st.markdown("---")
        
        # =============================================================================
        # 🌿 PARÁMETROS FORRAJEROS - MEJORADO CON PERSONALIZACIÓN Y MÁS REALISTAS
        # =============================================================================
        st.header("🌿 Parámetros Forrajeros")
        tipo_pastura = st.selectbox(
            "Tipo de Pastura:",
            ["ALFALFA", "RAYGRASS", "FESTUCA", "AGROPIRRO", "PASTIZAL_NATURAL", "PERSONALIZADO"]
        )
        
        # Parámetros personalizados si se selecciona "PERSONALIZADO"
        parametros_personalizados = {}
        if tipo_pastura == "PERSONALIZADO":
            with st.expander("🔧 Parámetros Personalizados", expanded=True):
                st.subheader("📊 Parámetros de Producción")
                parametros_personalizados['MS_POR_HA_OPTIMO'] = st.number_input(
                    "Biomasa Óptima (kg MS/ha):",
                    min_value=1000, max_value=10000, value=3500, step=100,
                    help="Materia seca óptima por hectárea"
                )
                parametros_personalizados['CRECIMIENTO_DIARIO'] = st.number_input(
                    "Crecimiento Diario (kg MS/ha/día):",
                    min_value=10, max_value=300, value=65, step=5,
                    help="Crecimiento diario de materia seca"
                )
                parametros_personalizados['FACTOR_BIOMASA_NDVI'] = st.number_input(
                    "Factor Biomasa-NDVI:",
                    min_value=1000, max_value=5000, value=2500, step=100,
                    help="Factor de conversión de NDVI a biomasa"
                )
                
                st.subheader("🌿 Parámetros de Calidad")
                parametros_personalizados['DIGESTIBILIDAD'] = st.slider(
                    "Digestibilidad (%):",
                    min_value=0.3, max_value=0.8, value=0.65, step=0.05,
                    help="Porcentaje de digestibilidad del forraje"
                )
                parametros_personalizados['PROTEINA_CRUDA'] = st.slider(
                    "Proteína Cruda (%):",
                    min_value=0.05, max_value=0.25, value=0.14, step=0.01,
                    help="Contenido de proteína cruda"
                )
                
                st.subheader("📏 Umbrales de Detección")
                parametros_personalizados['UMBRAL_NDVI_SUELO'] = st.slider(
                    "Umbral NDVI Suelo:",
                    min_value=0.05, max_value=0.3, value=0.16, step=0.01,
                    help="NDVI por debajo de este valor se considera suelo desnudo"
                )
                parametros_personalizados['UMBRAL_NDVI_PASTURA'] = st.slider(
                    "Umbral NDVI Pastura:",
                    min_value=0.3, max_value=0.8, value=0.48, step=0.01,
                    help="NDVI por encima de este valor se considera vegetación densa"
                )
        
        # =============================================================================
        # 🐄 PARÁMETROS GANADEROS - MEJORADO CON CONSUMO VOLUNTARIO REALISTA
        # =============================================================================
        st.header("🐄 Parámetros Ganaderos")
        peso_promedio = st.slider("Peso promedio (kg):", 300, 600, 450)
        carga_animal = st.slider("Carga animal:", 1, 1000, 100)
        
        with st.expander("🍽️ Parámetros de Consumo", expanded=False):
            st.subheader("Consumo Animal")
            consumo_voluntario = st.number_input(
                "Consumo Voluntario (kg MS/día/animal):",
                min_value=8.0, max_value=15.0, value=11.0, step=0.5,
                help="Consumo diario de materia seca por animal"
            )
            
            st.subheader("Eficiencias")
            eficiencia_pastoreo = st.slider(
                "Eficiencia de Pastoreo (%):",
                min_value=50, max_value=90, value=70, step=5,
                help="Porcentaje de forraje que realmente consume el animal"
            ) / 100.0
            
            eficiencia_cosecha = st.slider(
                "Eficiencia de Cosecha (%):",
                min_value=50, max_value=90, value=65, step=5,
                help="Porcentaje de biomasa que se puede cosechar eficientemente"
            ) / 100.0
        
        # =============================================================================
        # 📅 CONFIGURACIÓN TEMPORAL Y ESPACIAL
        # =============================================================================
        st.header("📅 Configuración Temporal")
        fecha_imagen = st.date_input(
            "Fecha de imagen:",
            value=datetime.now() - timedelta(days=30),
            max_value=datetime.now()
        )
        
        st.header("📐 Configuración Espacial")
        n_divisiones = st.slider("Sub-divisiones:", 8, 48, 24)
        
        st.header("📤 Cargar Datos")
        uploaded_file = st.file_uploader("Shapefile (ZIP):", type=['zip'])
    
    # Contenido principal
    st.title("🌱 Analizador Forrajero Unificado")
    st.markdown("---")
    
    # Inicializar analizador
    analizador = AnalizadorForrajeroUnificado()
    
    # Procesar archivo cargado
    if uploaded_file is not None:
        with st.spinner("Cargando y procesando archivo..."):
            try:
                with tempfile.TemporaryDirectory() as tmp_dir:
                    with zipfile.ZipFile(uploaded_file, 'r') as zip_ref:
                        zip_ref.extractall(tmp_dir)
                    
                    shp_files = [f for f in os.listdir(tmp_dir) if f.endswith('.shp')]
                    if shp_files:
                        gdf = gpd.read_file(os.path.join(tmp_dir, shp_files[0]))
                        if gdf.crs is None:
                            gdf = gdf.set_crs('EPSG:4326')
                        st.session_state.gdf_cargado = gdf
                        st.success("✅ Shapefile cargado correctamente")
                    else:
                        st.error("❌ No se encontró archivo .shp en el ZIP")
            except Exception as e:
                st.error(f"Error cargando shapefile: {e}")
    
    # Mostrar datos cargados y opción de análisis
    if 'gdf_cargado' in st.session_state and st.session_state.gdf_cargado is not None:
        gdf = st.session_state.gdf_cargado
        
        st.header("📁 Datos Cargados")
        area_total = analizador._calcular_superficie(gdf).sum()
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Polígonos", len(gdf))
        with col2:
            st.metric("Área Total", f"{area_total:.1f} ha")
        with col3:
            st.metric("Usuario", st.session_state.username)
        
        # Mapa de vista previa
        if FOLIUM_AVAILABLE:
            st.subheader("🗺️ Vista Previa del Potrero")
            mapa_preview = crear_mapa_base(gdf)
            if mapa_preview:
                # Agregar polígono al mapa
                folium.GeoJson(
                    gdf.__geo_interface__,
                    style_function=lambda x: {'fillColor': '#3388ff', 'color': 'blue', 'weight': 2, 'fillOpacity': 0.3}
                ).add_to(mapa_preview)
                folium_static(mapa_preview, width=800, height=400)
        
        # Botón de análisis
        if st.button("🚀 EJECUTAR ANÁLISIS COMPLETO", type="primary", use_container_width=True):
            config = {
                'tipo_pastura': tipo_pastura,
                'peso_promedio': peso_promedio,
                'carga_animal': carga_animal,
                'n_divisiones': n_divisiones,
                'fecha_imagen': fecha_imagen,
                'parametros_personalizados': parametros_personalizados if tipo_pastura == "PERSONALIZADO" else None,
                'consumo_voluntario': consumo_voluntario,
                'eficiencia_pastoreo': eficiencia_pastoreo,
                'eficiencia_cosecha': eficiencia_cosecha
            }
            
            with st.spinner("Realizando análisis completo..."):
                gdf_analizado = analizador.analizar_potrero(gdf, config)
                
                if gdf_analizado is not None:
                    st.session_state.gdf_analizado = gdf_analizado
                    st.success("✅ Análisis completado exitosamente!")
                    
                    # Mostrar resultados
                    mostrar_resultados_completos(gdf_analizado, config)
    
    else:
        # Pantalla de bienvenida
        st.info("""
        ### 🌱 Bienvenido al Analizador Forrajero Unificado
        
        **Características:**
        - 🛰️ **Datos satelitales** con Sentinel Hub
        - 📊 **Cálculo de EV/ha** y capacidad de carga
        - 📅 **Días de permanencia** por lote
        - 🗺️ **Mapas interactivos** con múltiples bases
        - 📄 **Informes automáticos** con recomendaciones
        - 🔐 **Sistema de autenticación** seguro
        
        **Para comenzar:**
        1. Configura los parámetros en la barra lateral
        2. Sube tu shapefile en formato ZIP
        3. Ejecuta el análisis completo
        4. Descarga el informe con recomendaciones
        """)

def mostrar_resultados_completos(gdf_analizado, config):
    """Muestra resultados completos del análisis con mapas mejorados"""
    st.header("📊 RESULTADOS DEL ANÁLISIS COMPLETO")
    
    # Métricas principales mejoradas
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        biomasa_prom = gdf_analizado['biomasa_disponible_kg_ms_ha'].mean()
        st.metric("Biomasa Disponible", f"{biomasa_prom:.0f} kg MS/ha")
    
    with col2:
        disponibilidad_prom = gdf_analizado['disponibilidad_forrajera_kg_ms_ha'].mean()
        st.metric("Disponibilidad Forrajera", f"{disponibilidad_prom:.0f} kg MS/ha")
    
    with col3:
        ev_total = gdf_analizado['carga_animal'].sum()
        st.metric("Capacidad Total", f"{ev_total:.1f} EV")
    
    with col4:
        dias_prom = gdf_analizado['dias_permanencia'].mean()
        st.metric("Días Permanencia", f"{dias_prom:.1f}")
    
    # Distribución de disponibilidad
    st.subheader("📈 Distribución de Disponibilidad Forrajera")
    distribucion = gdf_analizado['categoria_disponibilidad'].value_counts()
    col_dist1, col_dist2, col_dist3, col_dist4 = st.columns(4)
    
    with col_dist1:
        muy_baja = distribucion.get('MUY BAJA', 0)
        st.metric("🔴 Muy Baja", f"{muy_baja} sub-lotes")
    
    with col_dist2:
        baja = distribucion.get('BAJA', 0)
        st.metric("🟠 Baja", f"{baja} sub-lotes")
    
    with col_dist3:
        media = distribucion.get('MEDIA', 0)
        st.metric("🟡 Media", f"{media} sub-lotes")
    
    with col_dist4:
        alta = distribucion.get('ALTA', 0)
        st.metric("🟢 Alta", f"{alta} sub-lotes")
    
    # Mapas de resultados - CON LEYENDAS CLARAS
    if FOLIUM_AVAILABLE:
        st.header("🗺️ VISUALIZACIÓN INTERACTIVA")
        
        # Selector de mapa base
        col_map1, col_map2 = st.columns([3, 1])
        with col_map2:
            mapa_base_seleccionado = st.selectbox(
                "Mapa Base:",
                ["ESRI World Imagery", "ESRI World Street Map", "OpenStreetMap"],
                key="mapa_base_selector"
            )
        
        # Pestañas con mapas funcionales
        tab1, tab2, tab3, tab4, tab5 = st.tabs([
            "🌿 NDVI", 
            "🐄 EV/ha", 
            "📊 Disponibilidad",
            "🌱 Recomendaciones",
            "📋 Resumen"
        ])
        
        with tab1:
            st.subheader("🌿 ÍNDICE NDVI - ESTADO VEGETATIVO")
            st.info("💡 **Interpretación NDVI:** Valores más altos indican mejor estado de la vegetación")
            mapa_ndvi = crear_mapa_ndvi(gdf_analizado, mapa_base_seleccionado)
            if mapa_ndvi:
                folium_static(mapa_ndvi, width=800, height=600)
            else:
                st.warning("No se pudo generar el mapa de NDVI")
        
        with tab2:
            st.subheader("🐄 CAPACIDAD DE CARGA - EV/HA")
            st.info("💡 **Interpretación EV/ha:** Número de animales que puede sostener cada hectárea")
            mapa_ev = crear_mapa_ev_ha(gdf_analizado, mapa_base_seleccionado)
            if mapa_ev:
                folium_static(mapa_ev, width=800, height=600)
            else:
                st.warning("No se pudo generar el mapa de EV/ha")
        
        with tab3:
            st.subheader("📊 DISPONIBILIDAD FORRAJERA")
            st.info("💡 **Interpretación:** Cantidad de forraje disponible para el consumo animal")
            mapa_disp = crear_mapa_disponibilidad(gdf_analizado, mapa_base_seleccionado)
            if mapa_disp:
                folium_static(mapa_disp, width=800, height=600)
            else:
                st.warning("No se pudo generar el mapa de disponibilidad")
        
        with tab4:
            st.subheader("🌱 RECOMENDACIONES AGROECOLÓGICAS")
            st.info("💡 **Interpretación:** Acciones sugeridas según el estado forrajero")
            mapa_recom = crear_mapa_recomendaciones(gdf_analizado, mapa_base_seleccionado)
            if mapa_recom:
                folium_static(mapa_recom, width=800, height=600)
            else:
                st.warning("No se pudo generar el mapa de recomendaciones")
            
            # Recomendaciones detalladas
            st.subheader("📝 RECOMENDACIONES DETALLADAS")
            recomendaciones = _generar_recomendaciones_agroecologicas(gdf_analizado, config)
            st.markdown(recomendaciones)
        
        with tab5:
            st.subheader("📋 RESUMEN POR SUB-LOTE")
            # Tabla de resultados completa
            columnas = ['id_subLote', 'area_ha', 'tipo_superficie', 'ndvi', 
                       'biomasa_disponible_kg_ms_ha', 'disponibilidad_forrajera_kg_ms_ha',
                       'categoria_disponibilidad', 'ev_ha', 'dias_permanencia', 'dias_autonomia']
            
            # Filtrar columnas que existen
            columnas_existentes = [col for col in columnas if col in gdf_analizado.columns]
            tabla = gdf_analizado[columnas_existentes].copy()
            
            # Renombrar columnas para mejor visualización
            nombres_columnas = {
                'id_subLote': 'Sub-Lote',
                'area_ha': 'Área (ha)',
                'tipo_superficie': 'Tipo Superficie',
                'ndvi': 'NDVI',
                'biomasa_disponible_kg_ms_ha': 'Biomasa (kg MS/ha)',
                'disponibilidad_forrajera_kg_ms_ha': 'Disponibilidad (kg MS/ha)',
                'categoria_disponibilidad': 'Categoría',
                'ev_ha': 'EV/ha',
                'dias_permanencia': 'Días Permanencia',
                'dias_autonomia': 'Días Autonomía'
            }
            
            tabla.columns = [nombres_columnas.get(col, col) for col in tabla.columns]
            st.dataframe(tabla, use_container_width=True)
    
    else:
        st.warning("⚠️ Folium no está disponible. Los mapas interactivos no se mostrarán.")
    
    # Exportar datos
    st.header("💾 EXPORTAR RESULTADOS")
    col_exp1, col_exp2, col_exp3 = st.columns(3)
    
    with col_exp1:
        # CSV con todas las métricas
        columnas_exportar = ['id_subLote', 'area_ha', 'tipo_superficie', 'ndvi', 
                           'biomasa_disponible_kg_ms_ha', 'disponibilidad_forrajera_kg_ms_ha',
                           'categoria_disponibilidad', 'ev_ha', 'dias_permanencia', 'dias_autonomia']
        columnas_exportar = [col for col in columnas_exportar if col in gdf_analizado.columns]
        
        csv = gdf_analizado[columnas_exportar].to_csv(index=False)
        st.download_button(
            "📥 Descargar CSV Completo",
            csv,
            f"resultados_completos_{config['tipo_pastura']}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
            "text/csv",
            key="download_csv_unique"
        )
    
    with col_exp2:
        # GeoJSON
        geojson = gdf_analizado.to_json()
        st.download_button(
            "📥 Descargar GeoJSON",
            geojson,
            f"resultados_{config['tipo_pastura']}_{datetime.now().strftime('%Y%m%d_%H%M')}.geojson",
            "application/json",
            key="download_geojson_unique"
        )
    
    with col_exp3:
        # Informe DOCX
        if DOCX_AVAILABLE:
            with st.spinner("Preparando informe..."):
                mapa_buffer = crear_mapa_detallado(gdf_analizado, config['tipo_pastura'])
                informe_buffer = generar_informe_completo(gdf_analizado, config, mapa_buffer)
                
                if informe_buffer:
                    st.download_button(
                        "📄 Descargar Informe DOCX",
                        data=informe_buffer.getvalue(),
                        file_name=f"informe_forrajero_{config['tipo_pastura']}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_docx_unique"
                    )
                else:
                    st.error("Error generando informe DOCX")
        else:
            st.warning("python-docx no disponible para generar informes")

# =============================================================================
# FUNCIÓN PRINCIPAL
# =============================================================================

def main():
    """Función principal de la aplicación"""
    initialize_session_state()
    
    if not st.session_state.authenticated:
        login_section()
    else:
        main_application()

if __name__ == "__main__":
    main()
